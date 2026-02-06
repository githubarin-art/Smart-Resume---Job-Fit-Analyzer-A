"""
Debug script to inspect DOCX structure.
"""
from docx import Document
from docx.oxml.ns import qn

file_path = 'uploads/7a062df0-540b-4002-9a94-84ae5136488e.docx'
doc = Document(file_path)

print(f"Paragraphs: {len(doc.paragraphs)}")
print(f"Tables: {len(doc.tables)}")

print("\n--- Paragraphs ---")
for i, p in enumerate(doc.paragraphs[:20]):
    text = p.text[:80] if p.text else "(empty)"
    print(f"P{i}: {text}")

print("\n--- Tables ---")
for ti, table in enumerate(doc.tables):
    print(f"Table {ti}: {len(table.rows)} rows, {len(table.columns)} cols")
    for ri, row in enumerate(table.rows[:5]):
        for ci, cell in enumerate(row.cells):
            if cell.text.strip():
                print(f"  Cell[{ri},{ci}]: {cell.text[:50]}")

# Check for text boxes (shapes)
print("\n--- Looking for text boxes in document.xml ---")
body = doc._element.body
# Look for drawing/textbox elements
drawings = body.findall('.//' + qn('w:drawing'))
print(f"Drawings found: {len(drawings)}")

# Try to extract text from all XML
all_text = []
for elem in body.iter():
    if elem.text and elem.text.strip():
        all_text.append(elem.text.strip())
    if elem.tail and elem.tail.strip():
        all_text.append(elem.tail.strip())

print(f"\n--- All extracted text ({len(all_text)} parts) ---")
full_text = ' '.join(all_text)
print(full_text[:2000])
